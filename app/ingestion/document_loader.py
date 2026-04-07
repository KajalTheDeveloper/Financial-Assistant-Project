"""
Document Loader

Multi-format document loading with metadata extraction.
Supports PDF, TXT, CSV, DOCX, and web pages.
"""

import hashlib
from datetime import datetime
from pathlib import Path
from typing import BinaryIO, Optional, Union

try:
    from langchain_core.documents import Document
except Exception:
    # Minimal fallback Document dataclass for environments without langchain_core
    from dataclasses import dataclass

    @dataclass
    class Document:
        page_content: str
        metadata: dict

from app.core.config import settings
from app.core.exceptions import (
    DocumentParsingError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.core.logging import get_logger

logger = get_logger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".docx", ".html"}


class DocumentLoader:
    """
    Multi-format document loader with metadata extraction.
    
    Supports:
    - PDF files (with table extraction)
    - Text files (.txt, .md)
    - CSV files
    - Word documents (.docx)
    - HTML files
    """
    
    def __init__(self):
        self.supported_extensions = SUPPORTED_EXTENSIONS
    
    def load(
        self,
        file_path: Union[str, Path],
        file_obj: Optional[BinaryIO] = None,
        metadata: Optional[dict] = None
    ) -> list[Document]:
        """
        Load a document from file path or file object.
        
        Args:
            file_path: Path to the document
            file_obj: Optional file object (for uploaded files)
            metadata: Optional additional metadata
        
        Returns:
            List of Document objects with content and metadata
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise UnsupportedFileTypeError(
                extension,
                list(self.supported_extensions)
            )
        
        logger.info("Loading document", filename=file_path.name, extension=extension)
        
        try:
            # Load based on file type
            if extension == ".pdf":
                documents = self._load_pdf(file_path, file_obj)
            elif extension in {".txt", ".md"}:
                documents = self._load_text(file_path, file_obj)
            elif extension == ".csv":
                documents = self._load_csv(file_path, file_obj)
            elif extension == ".docx":
                documents = self._load_docx(file_path, file_obj)
            elif extension == ".html":
                documents = self._load_html(file_path, file_obj)
            else:
                raise UnsupportedFileTypeError(extension, list(self.supported_extensions))
            
            if not documents:
                raise EmptyDocumentError(file_path.name)
            
            # Add common metadata
            file_hash = self._compute_hash(file_path, file_obj)
            for doc in documents:
                doc.metadata.update({
                    "source_file": file_path.name,
                    "file_path": str(file_path),
                    "file_type": extension[1:],  # Remove dot
                    "file_hash": file_hash,
                    "ingested_at": datetime.now().isoformat(),
                })
                if metadata:
                    doc.metadata.update(metadata)
            
            logger.info(
                "Document loaded successfully",
                filename=file_path.name,
                num_pages=len(documents)
            )
            return documents
            
        except (UnsupportedFileTypeError, EmptyDocumentError):
            raise
        except Exception as e:
            logger.error("Failed to load document", filename=file_path.name, error=str(e))
            raise DocumentParsingError(file_path.name, str(e))
    
    def _load_pdf(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> list[Document]:
        """Load PDF document with table extraction."""
        import pdfplumber
        
        documents = []
        
        # Open from file object or path
        if file_obj:
            pdf = pdfplumber.open(file_obj)
        else:
            pdf = pdfplumber.open(file_path)
        
        try:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract text
                text = page.extract_text() or ""
                
                # Extract tables and convert to text
                tables = page.extract_tables()
                table_text = ""
                if tables:
                    for table in tables:
                        table_text += self._table_to_text(table) + "\n\n"
                
                # Combine text and tables
                content = text
                if table_text:
                    content += "\n\n[TABLES]\n" + table_text
                
                if content.strip():
                    documents.append(Document(
                        page_content=content,
                        metadata={
                            "page_number": page_num,
                            "total_pages": len(pdf.pages),
                            "has_tables": bool(tables),
                        }
                    ))
        finally:
            pdf.close()
        
        return documents
    
    def _load_text(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> list[Document]:
        """Load text file."""
        if file_obj:
            content = file_obj.read().decode("utf-8")
        else:
            content = file_path.read_text(encoding="utf-8")
        
        if not content.strip():
            return []
        
        return [Document(
            page_content=content,
            metadata={"page_number": 1, "total_pages": 1}
        )]
    
    def _load_csv(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> list[Document]:
        """Load CSV file as structured text."""
        import pandas as pd
        
        if file_obj:
            df = pd.read_csv(file_obj)
        else:
            df = pd.read_csv(file_path)
        
        # Convert to readable text format
        content = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n\n"
        content += f"Columns: {', '.join(df.columns)}\n\n"
        content += df.to_string(index=False)
        
        return [Document(
            page_content=content,
            metadata={
                "page_number": 1,
                "total_pages": 1,
                "num_rows": len(df),
                "num_columns": len(df.columns),
            }
        )]
    
    def _load_docx(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> list[Document]:
        """Load Word document."""
        from docx import Document as DocxDocument
        
        if file_obj:
            doc = DocxDocument(file_obj)
        else:
            doc = DocxDocument(file_path)
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        
        if not content.strip():
            return []
        
        return [Document(
            page_content=content,
            metadata={"page_number": 1, "total_pages": 1}
        )]
    
    def _load_html(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> list[Document]:
        """Load HTML file with content extraction."""
        from bs4 import BeautifulSoup
        
        if file_obj:
            html_content = file_obj.read().decode("utf-8")
        else:
            html_content = file_path.read_text(encoding="utf-8")
        
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        # Get text
        text = soup.get_text(separator="\n", strip=True)
        
        if not text.strip():
            return []
        
        return [Document(
            page_content=text,
            metadata={"page_number": 1, "total_pages": 1}
        )]
    
    def _table_to_text(self, table: list[list]) -> str:
        """Convert extracted table to readable text format."""
        if not table:
            return ""
        
        # Clean None values
        cleaned_table = [
            [str(cell) if cell else "" for cell in row]
            for row in table
        ]
        
        # Format as text table
        lines = []
        for row in cleaned_table:
            lines.append(" | ".join(row))
        
        return "\n".join(lines)
    
    def _compute_hash(
        self,
        file_path: Path,
        file_obj: Optional[BinaryIO] = None
    ) -> str:
        """Compute SHA256 hash for deduplication."""
        sha256 = hashlib.sha256()
        
        if file_obj:
            file_obj.seek(0)
            for chunk in iter(lambda: file_obj.read(8192), b""):
                sha256.update(chunk)
            file_obj.seek(0)
        else:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    sha256.update(chunk)
        
        return sha256.hexdigest()[:16]  # First 16 chars


def load_documents(
    source: Union[str, Path, list[Union[str, Path]]],
    metadata: Optional[dict] = None
) -> list[Document]:
    """
    Convenience function to load documents.
    
    Args:
        source: File path, directory path, or list of paths
        metadata: Optional metadata to add to all documents
    
    Returns:
        List of loaded Document objects
    """
    loader = DocumentLoader()
    all_documents = []
    
    if isinstance(source, (str, Path)):
        source = Path(source)
        if source.is_dir():
            # Load all supported files from directory
            for ext in SUPPORTED_EXTENSIONS:
                for file_path in source.glob(f"**/*{ext}"):
                    try:
                        docs = loader.load(file_path, metadata=metadata)
                        all_documents.extend(docs)
                    except Exception as e:
                        logger.warning(f"Skipping {file_path}: {e}")
        else:
            all_documents = loader.load(source, metadata=metadata)
    else:
        # List of paths
        for path in source:
            try:
                docs = loader.load(path, metadata=metadata)
                all_documents.extend(docs)
            except Exception as e:
                logger.warning(f"Skipping {path}: {e}")
    
    return all_documents
